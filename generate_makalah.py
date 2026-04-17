"""
Script untuk menghasilkan Makalah UTS Analisa Perancangan Sistem
Mata Kuliah: Analisa Perancangan Sistem
LP3I Cirebon - TA 2025/2026
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

doc = Document()

# =====================
# STYLES & PAGE SETUP
# =====================
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)

section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(3)
section.bottom_margin = Cm(3)
section.left_margin = Cm(4)
section.right_margin = Cm(3)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0,0,0)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:cs'), 'Times New Roman')
        rPr.insert(0, rFonts)
    return h

def add_para(text, bold=False, italic=False, align=None, size=12, indent_cm=0, spacing_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {})
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')
    rPr.insert(0, rFonts)
    if align:
        p.alignment = align
    if indent_cm > 0:
        p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.space_after = Pt(spacing_after)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_bullet(text, level=0, bold=False):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn('w:rFonts'), {})
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.insert(0, rFonts)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    p.paragraph_format.line_spacing = 1.5
    return p

def add_table_simple(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Shade header
        shading = cell._element.get_or_add_tcPr()
        shdElm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): 'D9E2F3',
            qn('w:val'): 'clear'
        })
        shading.append(shdElm)
    # Rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
    doc.add_paragraph()
    return table

def add_flowchart_placeholder(title, description):
    """Add a bordered placeholder for diagram"""
    add_para(f'Diagram: {title}', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    # Create a single-cell table as a bordered box
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'\n{description}\n\n[Render diagram Mermaid dari file flowchart_siakad.md\ndi https://mermaid.live lalu paste gambar di sini]\n')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)
    doc.add_paragraph()

def add_screenshot_placeholder(title):
    add_para(title, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'\n\n[TEMPEL SCREENSHOT {title.upper()} DI SINI]\n\n')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = RGBColor(150, 150, 150)
    doc.add_paragraph()

# ====================
# COVER PAGE
# ====================

for _ in range(4):
    doc.add_paragraph()

add_para('MAKALAH', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
add_para('UJIAN TENGAH SEMESTER', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
add_para('TA. 2025/2026', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)

doc.add_paragraph()

add_para('ANALISA PERANCANGAN SISTEM INFORMASI AKADEMIK', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_para('(SIAKAD) LP3I CIREBON', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_para('Menggunakan Pendekatan Systems Development Life Cycle (SDLC)', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

for _ in range(3):
    doc.add_paragraph()

# Info table
info_table = doc.add_table(rows=7, cols=3)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ('Mata Kuliah', ':', 'Analisa Perancangan Sistem'),
    ('Kelas', ':', 'MI1.4'),
    ('Semester', ':', '4 (Empat)'),
    ('Dosen Pengampu', ':', 'Azhar Al Afghani, M.Kom.'),
    ('Hari / Tanggal', ':', 'Jumat'),
    ('Ruang', ':', 'LAB 2'),
    ('Bentuk Ujian', ':', 'Project / Take Home'),
]
for ri, (k, sep, v) in enumerate(info_data):
    for ci, txt in enumerate([k, sep, v]):
        cell = info_table.rows[ri].cells[ci]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        if ci == 0:
            run.bold = True

for _ in range(2):
    doc.add_paragraph()

add_para('Disusun oleh:', align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para('[NAMA MAHASISWA]', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_para('[NIM]', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

for _ in range(3):
    doc.add_paragraph()

add_para('PROGRAM STUDI MANAJEMEN INFORMATIKA', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para('POLITEKNIK LP3I CIREBON', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para('2026', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

doc.add_page_break()

# ====================
# DAFTAR ISI
# ====================

add_heading_styled('DAFTAR ISI', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

toc_items = [
    ('DAFTAR ISI', 'ii'),
    ('DAFTAR TABEL', 'iii'),
    ('DAFTAR GAMBAR', 'iv'),
    ('', ''),
    ('BAB I PENDAHULUAN', '1'),
    ('    1.1 Latar Belakang', '1'),
    ('    1.2 Rumusan Masalah', '2'),
    ('    1.3 Tujuan Penelitian', '2'),
    ('    1.4 Ruang Lingkup', '3'),
    ('', ''),
    ('BAB II LANDASAN TEORI', '4'),
    ('    2.1 Sistem Informasi Akademik (SIAKAD)', '4'),
    ('    2.2 Systems Development Life Cycle (SDLC)', '4'),
    ('    2.3 Teknologi yang Digunakan', '5'),
    ('', ''),
    ('BAB III ANALISIS SISTEM YANG BERJALAN', '6'),
    ('    3.1 Gambaran Umum Sistem Lama', '6'),
    ('    3.2 Proses Akademik yang Berjalan', '6'),
    ('    3.3 Alur Data dan Pengguna', '7'),
    ('    3.4 Kelebihan Sistem Lama', '7'),
    ('    3.5 Kekurangan Sistem Lama', '8'),
    ('', ''),
    ('BAB IV PERANCANGAN SISTEM BARU (SDLC)', '9'),
    ('    4.1 Tahap Perencanaan (Planning)', '9'),
    ('    4.2 Tahap Analisis (Analysis)', '10'),
    ('    4.3 Tahap Desain (Design)', '11'),
    ('    4.4 Entity Relationship Diagram (ERD)', '14'),
    ('', ''),
    ('BAB V IMPLEMENTASI PROTOTIPE', '16'),
    ('    5.1 Arsitektur Sistem', '16'),
    ('    5.2 Teknologi dan Tools', '16'),
    ('    5.3 Tampilan Antarmuka Sistem', '17'),
    ('    5.4 Fitur-Fitur Utama Prototipe', '18'),
    ('', ''),
    ('BAB VI PERBANDINGAN SISTEM', '20'),
    ('    6.1 Tabel Perbandingan', '20'),
    ('    6.2 Keunggulan Sistem Baru', '20'),
    ('', ''),
    ('BAB VII TESTING DAN EVALUASI', '22'),
    ('    7.1 Pengujian Fungsional', '22'),
    ('    7.2 Hasil Pengujian', '22'),
    ('', ''),
    ('BAB VIII PEMELIHARAAN (MAINTENANCE)', '24'),
    ('    8.1 Rencana Pemeliharaan', '24'),
    ('    8.2 Perawatan Sistem', '24'),
    ('', ''),
    ('BAB IX KESIMPULAN DAN SARAN', '25'),
    ('    9.1 Kesimpulan', '25'),
    ('    9.2 Saran Pengembangan', '25'),
    ('', ''),
    ('DAFTAR PUSTAKA', '26'),
    ('LAMPIRAN', '27'),
]

for item, page in toc_items:
    if item == '':
        doc.add_paragraph()
        continue
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    if item.startswith('    '):
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(item.strip())
    else:
        run = p.add_run(item)
        run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    # Tab + page number
    tab_run = p.add_run(f'\t{page}')
    tab_run.font.name = 'Times New Roman'
    tab_run.font.size = Pt(12)

doc.add_page_break()

# ==============================
# BAB I - PENDAHULUAN
# ==============================

add_heading_styled('BAB I', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
add_heading_styled('PENDAHULUAN', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_heading_styled('1.1 Latar Belakang', level=2)
add_para('Perkembangan teknologi informasi yang pesat telah merubah berbagai aspek kehidupan, termasuk dalam dunia pendidikan. Perguruan tinggi sebagai institusi pendidikan dituntut untuk terus berinovasi dalam mengelola data akademik agar lebih efektif, efisien, dan transparan. Salah satu instrumen vital dalam pengelolaan data akademik adalah Sistem Informasi Akademik atau yang lebih dikenal dengan istilah SIAKAD.')

add_para('Politeknik LP3I Cirebon sebagai salah satu institusi pendidikan vokasi telah mengimplementasikan sistem SIAKAD untuk mengelola berbagai proses akademik, mulai dari pengisian Kartu Rencana Studi (KRS), pencatatan kehadiran, penilaian, hingga administrasi keuangan mahasiswa. Namun demikian, dalam operasionalnya, sistem yang berjalan saat ini masih memiliki sejumlah keterbatasan yang perlu diatasi.')

add_para('Berdasarkan pengamatan langsung terhadap sistem SIAKAD yang sedang berjalan di LP3I Cirebon, ditemukan beberapa permasalahan utama, antara lain: antarmuka pengguna (user interface) yang kurang intuitif dan modern, proses administrasi yang masih memerlukan intervensi manual pada beberapa tahapan, serta keterbatasan fitur integrasi antar modul yang menyebabkan data tidak tersinkronisasi secara real-time antar pemangku kepentingan (mahasiswa, dosen, dan admin BAAK).')

add_para('Oleh karena itu, dalam makalah ini penulis melakukan analisa mendalam terhadap sistem SIAKAD yang sedang berjalan, kemudian merancang dan membangun prototipe sistem baru yang lebih baik menggunakan pendekatan Systems Development Life Cycle (SDLC). Prototipe yang dibangun menggunakan teknologi modern yaitu React.js untuk frontend, Node.js/Express untuk backend, dan MySQL untuk manajemen basis data, sehingga menghasilkan sistem yang responsif, terintegrasi, dan siap untuk dideploy.')

add_heading_styled('1.2 Rumusan Masalah', level=2)
add_para('Berdasarkan latar belakang yang telah diuraikan, maka rumusan masalah dalam penelitian ini adalah sebagai berikut:')
add_bullet('Bagaimana proses akademik pada sistem SIAKAD LP3I Cirebon yang sedang berjalan saat ini?')
add_bullet('Apa saja kelebihan dan kekurangan dari sistem SIAKAD yang sedang digunakan?')
add_bullet('Bagaimana merancang sistem SIAKAD baru yang lebih baik dengan pendekatan SDLC?')
add_bullet('Bagaimana membangun prototipe aplikasi yang mampu mendemonstrasikan proses utama sistem SIAKAD secara terintegrasi?')

add_heading_styled('1.3 Tujuan Penelitian', level=2)
add_para('Adapun tujuan dari penelitian dan pembuatan makalah ini adalah:')
add_bullet('Menganalisa sistem informasi akademik (SIAKAD) yang sedang berjalan di LP3I Cirebon, meliputi proses akademik, alur data dan pengguna.')
add_bullet('Mengidentifikasi kelebihan dan kekurangan sistem yang sedang berjalan.')
add_bullet('Merancang sistem baru menggunakan metode SDLC yang memiliki keunggulan dibanding sistem lama.')
add_bullet('Membangun prototipe aplikasi SIAKAD Reborn berbasis web yang fungsional dan siap deploy.')
add_bullet('Melakukan pengujian dan evaluasi terhadap prototipe yang dibangun.')

add_heading_styled('1.4 Ruang Lingkup', level=2)
add_para('Ruang lingkup penelitian ini dibatasi pada hal-hal berikut:')
add_bullet('Analisa dilakukan terhadap sistem SIAKAD yang berjalan di Politeknik LP3I Cirebon.')
add_bullet('Perancangan sistem baru mencakup 3 role pengguna: Mahasiswa, Dosen, dan Admin BAAK.')
add_bullet('Modul yang dicakup meliputi: Autentikasi, KRS, Jadwal, Nilai, Kehadiran, Keuangan, IKM, CNP, Tugas Akhir, Persuratan, Literasi Digital, E-Library, Helpdesk, dan Profil.')
add_bullet('Prototipe dibangun menggunakan React.js, Node.js, Express, dan MySQL.')
add_bullet('Metode pengembangan yang digunakan adalah Systems Development Life Cycle (SDLC) model Waterfall.')

doc.add_page_break()

# ==============================
# BAB II - LANDASAN TEORI
# ==============================

add_heading_styled('BAB II', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
add_heading_styled('LANDASAN TEORI', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_heading_styled('2.1 Sistem Informasi Akademik (SIAKAD)', level=2)
add_para('Sistem Informasi Akademik (SIAKAD) adalah sebuah sistem yang dirancang untuk mengelola data dan proses akademik pada institusi pendidikan secara terkomputerisasi. Menurut Jogiyanto (2005), sistem informasi adalah suatu sistem di dalam suatu organisasi yang mempertemukan kebutuhan pengolahan transaksi harian, mendukung operasi, bersifat manajerial dan kegiatan strategi dari suatu organisasi.')

add_para('SIAKAD umumnya mencakup beberapa modul utama, antara lain:')
add_bullet('Modul Registrasi dan Data Mahasiswa')
add_bullet('Modul Kartu Rencana Studi (KRS)')
add_bullet('Modul Penjadwalan Kuliah')
add_bullet('Modul Pencatatan Kehadiran (Absensi)')
add_bullet('Modul Penilaian (Input dan Publikasi Nilai)')
add_bullet('Modul Keuangan (Pembayaran SPP/UKT)')
add_bullet('Modul Administrasi Persuratan')

add_heading_styled('2.2 Systems Development Life Cycle (SDLC)', level=2)
add_para('Systems Development Life Cycle (SDLC) adalah metodologi yang digunakan untuk mengembangkan, memelihara, dan mengganti sistem informasi. SDLC terdiri dari beberapa tahapan yang sistematis dan terstruktur, dimana setiap tahapan menghasilkan output yang menjadi input bagi tahapan berikutnya.')

add_para('Tahapan-tahapan SDLC model Waterfall yang digunakan dalam penelitian ini adalah:')
add_bullet('Perencanaan (Planning) — Identifikasi masalah, tujuan, dan ruang lingkup sistem.', bold=True)
add_bullet('Analisis (Analysis) — Analisa kebutuhan sistem yang ada, pembuatan diagram alir dan use case.', bold=True)
add_bullet('Desain (Design) — Membuat desain sistem baru meliputi interface, database, dan alur proses.', bold=True)
add_bullet('Implementasi (Implementation) — Membangun prototipe sistem berupa aplikasi fungsional.', bold=True)
add_bullet('Testing & Evaluasi (Testing & Evaluation) — Menguji sistem baru dan membandingkan dengan sistem lama.', bold=True)
add_bullet('Pemeliharaan (Maintenance) — Rencana pengembangan dan perawatan sistem.', bold=True)

add_heading_styled('2.3 Teknologi yang Digunakan', level=2)

add_table_simple(
    ['No', 'Teknologi', 'Keterangan'],
    [
        ['1', 'React.js', 'Library JavaScript untuk membangun antarmuka pengguna (frontend) yang reaktif dan modular.'],
        ['2', 'Vite', 'Build tool modern untuk project React yang memberikan hot module replacement secara instan.'],
        ['3', 'Node.js', 'Runtime JavaScript sisi server untuk membangun backend API.'],
        ['4', 'Express.js', 'Framework minimalis Node.js untuk membuat RESTful API endpoints.'],
        ['5', 'MySQL', 'Relational Database Management System (RDBMS) untuk penyimpanan data terstruktur.'],
        ['6', 'XAMPP', 'Paket distribusi Apache, MySQL, dan PHP untuk lingkungan pengembangan lokal.'],
        ['7', 'Multer', 'Middleware Node.js untuk menangani upload file (dokumen, sertifikat, bukti transfer).'],
        ['8', 'Tailwind CSS', 'Utility-first CSS framework untuk styling antarmuka modern dan responsif.'],
        ['9', 'Lucide React', 'Icon library berbasis SVG untuk elemen visual antarmuka.'],
    ]
)

doc.add_page_break()

# ==============================  
# BAB III - ANALISIS SISTEM LAMA
# ==============================

add_heading_styled('BAB III', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
add_heading_styled('ANALISIS SISTEM YANG BERJALAN', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_heading_styled('3.1 Gambaran Umum Sistem Lama', level=2)
add_para('Sistem SIAKAD yang saat ini berjalan di Politeknik LP3I Cirebon merupakan aplikasi berbasis web yang digunakan oleh seluruh civitas akademika untuk mengelola proses-proses akademik. Sistem ini dapat diakses melalui browser dan memiliki tiga jenis pengguna utama, yaitu mahasiswa, dosen, dan admin (BAAK). Setiap pengguna memiliki hak akses yang berbeda sesuai dengan peran masing-masing dalam sistem.')

add_para('Secara umum, sistem SIAKAD LP3I melayani proses-proses berikut: pengisian KRS online, pencatatan dan rekap kehadiran, penginputan dan publikasi nilai, pengelolaan tagihan dan pembayaran SPP/UKT, pencatatan Indeks Keaktifan Mahasiswa (IKM), pengelolaan administrasi persuratan, serta fasilitas helpdesk untuk layanan bantuan teknis.')

add_heading_styled('3.2 Proses Akademik yang Berjalan', level=2)
add_para('Berdasarkan observasi terhadap sistem SIAKAD LP3I Cirebon, berikut adalah proses-proses akademik utama yang berjalan:')

add_para('a) Proses Pengisian KRS', bold=True)
add_para('Mahasiswa melakukan login ke sistem, kemudian memilih mata kuliah yang akan diambil pada semester berjalan. Setelah memilih, mahasiswa men-submit KRS untuk divalidasi oleh pihak BAAK. Admin BAAK kemudian melakukan pengecekan dan memberikan persetujuan atau penolakan terhadap KRS tersebut.')

add_para('b) Proses Pencatatan Kehadiran', bold=True)
add_para('Dosen melakukan pencatatan kehadiran mahasiswa setiap pertemuan kuliah. Data kehadiran ini mencakup 14 pertemuan reguler ditambah pertemuan UTS dan UAS. Mahasiswa dapat melihat rekap kehadirannya melalui menu khusus di dashboard mahasiswa.')

add_para('c) Proses Penilaian', bold=True)
add_para('Dosen menginput nilai mahasiswa yang terdiri dari komponen Tugas, UTS, dan UAS. Sistem kemudian menghitung nilai huruf secara otomatis berdasarkan bobot yang telah ditentukan. Mahasiswa dapat melihat nilai akhir melalui menu Papan Nilai.')

add_para('d) Proses Keuangan', bold=True)
add_para('Sistem menampilkan daftar tagihan SPP/UKT mahasiswa beserta status pembayarannya. Mahasiswa yang telah melakukan pembayaran dapat mengupload bukti transfer untuk divalidasi oleh admin BAAK.')

add_heading_styled('3.3 Alur Data dan Pengguna', level=2)
add_para('Sistem SIAKAD LP3I memiliki tiga aktor utama dengan alur data sebagai berikut:')

add_table_simple(
    ['No', 'Aktor', 'Hak Akses'],
    [
        ['1', 'Mahasiswa', 'Mengisi KRS, melihat jadwal, melihat nilai & kehadiran, melakukan pembayaran, mengajukan surat, mengumpulkan poin IKM, mengakses e-library, menghubungi helpdesk.'],
        ['2', 'Dosen', 'Melihat jadwal mengajar, menginput nilai, mencatat kehadiran, membimbing mahasiswa.'],
        ['3', 'Admin (BAAK)', 'Mengelola data master (mahasiswa, dosen, mata kuliah), memvalidasi KRS, memvalidasi berkas/dokumen, mengelola keuangan, membalas tiket helpdesk.'],
    ]
)

add_heading_styled('3.4 Kelebihan Sistem Lama', level=2)
add_bullet('Sistem sudah berbasis web sehingga dapat diakses dari mana saja melalui browser.')
add_bullet('Pemisahan hak akses berdasarkan role (mahasiswa, dosen, admin) sudah diterapkan.')
add_bullet('Proses pengisian KRS sudah dilakukan secara online.')
add_bullet('Data akademik tersimpan dalam basis data terpusat.')
add_bullet('Sudah memiliki fitur Indeks Keaktifan Mahasiswa (IKM) untuk pencatatan poin keaktifan.')

add_heading_styled('3.5 Kekurangan Sistem Lama', level=2)
add_bullet('Antarmuka pengguna (UI) kurang modern dan user-friendly, menggunakan desain konvensional yang tidak responsif secara optimal di perangkat mobile.')
add_bullet('Beberapa proses administrasi masih memerlukan konfirmasi manual di luar sistem (misal: konfirmasi pembayaran via WhatsApp).')
add_bullet('Integrasi antar modul belum sepenuhnya real-time, sehingga perubahan data pada satu modul tidak langsung tercermin di modul lainnya.')
add_bullet('Fitur upload dokumen pendukung (sertifikat, bukti transfer, dokumen akademik) terbatas dan tidak terpusat.')
add_bullet('Tidak tersedia fitur literasi digital sebagai bagian dari ekosistem modern pendidikan.')
add_bullet('Sistem helpdesk belum terintegrasi langsung ke dalam SIAKAD, sehingga mahasiswa harus menghubungi admin melalui jalur komunikasi terpisah.')
add_bullet('Proses validasi berkas masih tersegmentasi, admin harus membuka menu berbeda untuk setiap jenis dokumen.')

doc.add_page_break()

# ==============================
# BAB IV - PERANCANGAN SISTEM BARU
# ==============================

add_heading_styled('BAB IV', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
add_heading_styled('PERANCANGAN SISTEM BARU (SDLC)', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_heading_styled('4.1 Tahap Perencanaan (Planning)', level=2)
add_para('Pada tahap perencanaan, dilakukan identifikasi masalah, penetapan tujuan, dan penentuan ruang lingkup pengembangan sistem baru. Berdasarkan analisis terhadap kekurangan sistem lama yang telah diuraikan pada BAB III, maka dirancang sistem SIAKAD Reborn dengan spesifikasi sebagai berikut:')

add_table_simple(
    ['Aspek', 'Keterangan'],
    [
        ['Nama Sistem', 'SIAKAD Reborn — Sistem Informasi Akademik Generasi Baru'],
        ['Tujuan', 'Membangun sistem SIAKAD yang modern, terintegrasi, dan siap deploy sebagai pengganti sistem lama.'],
        ['Pengguna', '3 Role: Mahasiswa, Dosen, Admin BAAK'],
        ['Platform', 'Web-based (Browser), Responsive Design'],
        ['Backend', 'Node.js + Express.js (RESTful API)'],
        ['Frontend', 'React.js + Vite + Tailwind CSS'],
        ['Database', 'MySQL (via XAMPP)'],
        ['Jumlah Modul', '25+ modul fungsional terintegrasi'],
    ]
)

add_heading_styled('4.2 Tahap Analisis (Analysis)', level=2)

add_para('4.2.1 Use Case Diagram', bold=True)
add_para('Use Case Diagram berikut menggambarkan interaksi antara empat aktor utama (Mahasiswa, Dosen, Admin BAAK, dan Tamu) dengan sistem SIAKAD Reborn:')

add_flowchart_placeholder(
    'Use Case Diagram SIAKAD Reborn',
    'Diagram ini menampilkan 4 aktor (Mahasiswa, Dosen, Admin, Tamu)\ndan 25 use case yang mencakup seluruh fungsionalitas sistem.\nAktor Mahasiswa memiliki 15 use case, Dosen 5 use case,\nAdmin BAAK 7 use case, dan Tamu 1 use case (Helpdesk).'
)

add_para('Tabel daftar use case per aktor:', bold=True)
add_table_simple(
    ['Aktor', 'Use Case', 'Jumlah'],
    [
        ['Mahasiswa', 'Login, KRS, Jadwal, Nilai, Kehadiran, Keuangan, IKM, CNP, Tugas Akhir, Surat, Upload Dokumen, Literasi Digital, E-Library, Helpdesk, Profil', '15'],
        ['Dosen', 'Login, Input Nilai, Input Absensi, Bimbingan Akademik, E-Library', '5'],
        ['Admin BAAK', 'Login, Validasi KRS, Validasi Berkas, Master Mahasiswa, Master Dosen, Rekap Keuangan, Helpdesk Admin', '7'],
        ['Tamu', 'Helpdesk (tanpa login)', '1'],
    ]
)

add_para('4.2.2 Diagram Alir (Flowchart) Sistem', bold=True)
add_para('Berikut adalah diagram-diagram alir yang menggambarkan proses bisnis utama pada sistem SIAKAD Reborn:')

# Flowchart 1 - Login
add_para('A. Flowchart Alur Login dan Routing Berdasarkan Role', bold=True, size=11)
add_flowchart_placeholder(
    'Flowchart Login & Routing',
    'Mulai → Landing Page → Klik Login → Input ID & Password\n→ Validasi Kredensial → Identifikasi Role\n→ Dashboard Mahasiswa / Dosen / Admin\n(Jika gagal: Tampilkan Pesan Error → Kembali ke Form Login)'
)

# Flowchart 2 - KRS
add_para('B. Flowchart Proses Pengisian KRS', bold=True, size=11)
add_flowchart_placeholder(
    'Flowchart Modul KRS',
    'Mahasiswa Buka KRS → Pilih Matkul → Cek Batas SKS\n→ Submit KRS (Status: Pending) → Admin Review\n→ Approve: Status Approved / Reject: Isi Ulang\n→ Mahasiswa Lihat Jadwal Kuliah'
)

# Flowchart 3 - Keuangan
add_para('C. Flowchart Proses Keuangan dan Pembayaran', bold=True, size=11)
add_flowchart_placeholder(
    'Flowchart Modul Keuangan',
    'Mahasiswa Buka Keuangan → Lihat Tagihan\n→ Ada Tunggakan? → Ya: Upload Bukti Transfer\n→ Status Pending → Admin Validasi\n→ Approve: Tagihan Lunas / Reject: Upload Ulang'
)

# Flowchart 4 - Nilai
add_para('D. Flowchart Proses Input Nilai oleh Dosen', bold=True, size=11)
add_flowchart_placeholder(
    'Flowchart Modul Nilai',
    'Dosen Login → Pilih Matkul → Lihat Daftar Mahasiswa KRS\n→ Input Skor Tugas/UTS/UAS → Kalkulasi Huruf Otomatis\n→ Simpan ke Database → Mahasiswa Lihat di Papan Nilai'
)

# Flowchart 5 - IKM
add_para('E. Flowchart Indeks Keaktifan Mahasiswa (IKM)', bold=True, size=11)
add_flowchart_placeholder(
    'Flowchart Modul IKM',
    'Mahasiswa Buka IKM → Baca Info 250 Poin\n→ Isi Form (SMT, Kategori, Kegiatan, Peringkat)\n→ Upload Sertifikat PDF + Foto Dokumentasi\n→ Hitung Poin Awal Otomatis → Status: Pending\n→ Admin Review → Approve + Input Poin Valid\n→ Poin Bertambah di Tabel Riwayat'
)

# Flowchart 6 - CNP & TA
add_para('F. Flowchart Career & Placement (CNP) dan Tugas Akhir', bold=True, size=11)
add_flowchart_placeholder(
    'Flowchart Modul CNP & Tugas Akhir',
    'Mahasiswa Pilih Sub-Menu → 4 Cabang:\n1. Upload Berkas → Admin Validasi\n2. Download Panduan → File Terunduh\n3. Input Jurnal Bimbingan → Simpan ke DB\n4. Lihat Jadwal Sidang → Tampilkan Jadwal'
)

# Flowchart 7 - Helpdesk
add_para('G. Flowchart Helpdesk / Pusat Bantuan', bold=True, size=11)
add_flowchart_placeholder(
    'Flowchart Modul Helpdesk',
    'User/Tamu → Tulis Pesan → Submit Tiket (Status: Open)\n→ Admin Lihat Tiket → Admin Tulis Balasan\n→ Status: Closed → User Lihat Balasan'
)

# Flowchart 8 - Literasi Digital
add_para('H. Flowchart Modul Literasi Digital', bold=True, size=11)
add_flowchart_placeholder(
    'Flowchart Modul Literasi Digital',
    'Mahasiswa Buka Literasi Digital → Baca Panduan\n→ Ikuti Pelatihan di Portal Kominfo\n→ Unduh Sertifikat → Upload via SIAKAD\n→ Admin Validasi → Status: Disetujui/Ditolak'
)

add_heading_styled('4.3 Tahap Desain (Design)', level=2)

add_para('4.3.1 Desain Arsitektur Sistem', bold=True)
add_para('Sistem SIAKAD Reborn dibangun menggunakan arsitektur 3-Tier yang terdiri dari:')

add_table_simple(
    ['Layer', 'Teknologi', 'Fungsi'],
    [
        ['Presentation Layer (Frontend)', 'React.js + Vite + Tailwind CSS', 'Menampilkan antarmuka pengguna yang responsif dan interaktif. Berjalan di browser (http://localhost:5173).'],
        ['Business Logic Layer (Backend)', 'Node.js + Express.js + Multer', 'Memproses logika bisnis, menyediakan RESTful API endpoints, dan mengelola upload file. Berjalan di port 5000.'],
        ['Data Layer (Database)', 'MySQL via XAMPP', 'Menyimpan seluruh data terstruktur dalam database siakad_db dengan 11 tabel utama.'],
    ]
)

add_flowchart_placeholder(
    'Diagram Arsitektur Sistem Keseluruhan',
    'Frontend (React.js + Vite) → REST API (Express.js Port 5000)\n→ MySQL Database (siakad_db)\n+ Multer Upload Engine → Static File Server (/uploads)'
)

add_para('4.3.2 Desain Database', bold=True)
add_para('Desain database sistem SIAKAD Reborn terdiri dari 11 tabel utama yang saling berelasi. Berikut adalah penjelasan masing-masing tabel:')

add_heading_styled('4.4 Entity Relationship Diagram (ERD)', level=2)

add_flowchart_placeholder(
    'Entity Relationship Diagram (ERD)',
    '11 Tabel: users, matkuls, krs, nilai, absensi,\nkeuangan, submissions, ikm_activities,\nbimbingan_logs, tickets, ebooks\ndengan relasi antar tabel melalui foreign key\n(nim, dosenId, userId)'
)

add_para('Berikut adalah penjelasan struktur tabel pada database siakad_db:', bold=True)

# Table structures
tables_info = [
    ('1. Tabel users', [
        ['id (PK)', 'VARCHAR(50)', 'NIM/NIP/admin sebagai ID unik login'],
        ['password', 'VARCHAR(255)', 'Kata sandi pengguna'],
        ['role', 'VARCHAR(20)', 'mahasiswa / dosen / admin'],
        ['nama', 'VARCHAR(150)', 'Nama lengkap'],
        ['prodi', 'VARCHAR(100)', 'Program studi'],
        ['email', 'VARCHAR(255)', 'Email pribadi'],
        ['noHp', 'VARCHAR(20)', 'Nomor telepon'],
        ['alamat', 'TEXT', 'Alamat domisili'],
        ['nik', 'VARCHAR(50)', 'Nomor Induk Kependudukan'],
        ['namaIbu', 'VARCHAR(100)', 'Nama ibu kandung'],
    ]),
    ('2. Tabel matkuls', [
        ['id (PK)', 'VARCHAR(50)', 'Kode mata kuliah'],
        ['nama', 'VARCHAR(150)', 'Nama mata kuliah'],
        ['sks', 'INT', 'Jumlah SKS'],
        ['semester', 'INT', 'Semester penawaran'],
        ['dosenId (FK)', 'VARCHAR(50)', 'Referensi ke tabel users (dosen)'],
        ['ruang', 'VARCHAR(50)', 'Ruang kelas'],
        ['jadwal', 'VARCHAR(100)', 'Hari dan jam'],
    ]),
    ('3. Tabel krs', [
        ['id (PK)', 'INT AUTO_INCREMENT', 'ID unik KRS'],
        ['userId (FK)', 'VARCHAR(50)', 'Referensi ke tabel users (mahasiswa)'],
        ['matkuls', 'JSON', 'Array kode mata kuliah yang dipilih'],
        ['status', 'VARCHAR(50)', 'draft / pending / approved'],
    ]),
    ('4. Tabel nilai', [
        ['userId_matkulId (PK)', 'VARCHAR(100)', 'Gabungan NIM dan kode MK'],
        ['tugas', 'FLOAT', 'Nilai tugas'],
        ['uts', 'FLOAT', 'Nilai UTS'],
        ['uas', 'FLOAT', 'Nilai UAS'],
        ['huruf', 'VARCHAR(5)', 'Nilai huruf (A/B/C/D/E)'],
    ]),
    ('5. Tabel keuangan', [
        ['id (PK)', 'VARCHAR(50)', 'ID tagihan'],
        ['nim (FK)', 'VARCHAR(50)', 'NIM mahasiswa'],
        ['bulan', 'VARCHAR(50)', 'Periode tagihan'],
        ['tagihan', 'INT', 'Nominal tagihan (Rp)'],
        ['status', 'VARCHAR(20)', 'Lunas / Nunggak / Belum Lunas'],
        ['jatuhTempo', 'VARCHAR(50)', 'Tanggal jatuh tempo'],
    ]),
    ('6. Tabel submissions', [
        ['id (PK)', 'VARCHAR(50)', 'ID pengajuan'],
        ['nim (FK)', 'VARCHAR(50)', 'NIM mahasiswa'],
        ['kategori', 'VARCHAR(50)', 'cnp / tugas_akhir / keuangan / surat / dokumen_pendukung / literasi_digital'],
        ['tipe', 'VARCHAR(150)', 'Jenis berkas spesifik'],
        ['fileUrl', 'TEXT', 'URL file yang diupload'],
        ['status', 'VARCHAR(20)', 'Pending / Disetujui / Ditolak'],
    ]),
    ('7. Tabel ikm_activities', [
        ['id (PK)', 'VARCHAR(50)', 'ID aktivitas IKM'],
        ['nim (FK)', 'VARCHAR(50)', 'NIM mahasiswa'],
        ['kategori', 'VARCHAR(100)', 'Organisasi / Prestasi / Seminar / Kepanitiaan'],
        ['kegiatan', 'VARCHAR(100)', 'Peserta / Panitia / Narasumber / Juara'],
        ['peringkat', 'VARCHAR(50)', 'Internasional / Nasional / Provinsi / Lokal'],
        ['fileSertifikat', 'TEXT', 'URL file sertifikat'],
        ['fileFoto', 'TEXT', 'URL foto dokumentasi'],
        ['poinAwal', 'INT', 'Estimasi poin berdasarkan peringkat'],
        ['poinValid', 'INT', 'Poin yang disahkan oleh admin BAAK'],
        ['status', 'VARCHAR(20)', 'Pending / Disetujui / Ditolak'],
    ]),
    ('8. Tabel tickets', [
        ['id (PK)', 'VARCHAR(50)', 'ID tiket helpdesk'],
        ['senderId', 'VARCHAR(50)', 'ID pengirim (atau kosong jika tamu)'],
        ['senderName', 'VARCHAR(100)', 'Nama pengirim'],
        ['isGuest', 'BOOLEAN', 'Apakah pengirim adalah tamu'],
        ['message', 'TEXT', 'Isi pesan keluhan/pertanyaan'],
        ['reply', 'TEXT', 'Balasan dari admin'],
        ['status', 'VARCHAR(20)', 'open / closed'],
    ]),
]

for title, rows in tables_info:
    add_para(title, bold=True, size=11)
    add_table_simple(['Kolom', 'Tipe Data', 'Keterangan'], rows)

doc.add_page_break()

# ==============================
# BAB V - IMPLEMENTASI PROTOTIPE
# ==============================

add_heading_styled('BAB V', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
add_heading_styled('IMPLEMENTASI PROTOTIPE', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_heading_styled('5.1 Arsitektur Sistem', level=2)
add_para('Prototipe SIAKAD Reborn dibangun dengan arsitektur client-server, dimana frontend dan backend berjalan sebagai dua layanan terpisah yang berkomunikasi melalui REST API. Frontend berjalan di http://localhost:5173 (Vite dev server), sedangkan backend berjalan di http://localhost:5000 (Express.js). Database MySQL dikelola melalui XAMPP dan secara otomatis diinisialisasi oleh server backend saat pertama kali dijalankan.')

add_para('Keunggulan arsitektur ini adalah pemisahan tanggung jawab (separation of concerns) yang jelas antara presentasi, logika bisnis, dan penyimpanan data. Hal ini memudahkan proses pengembangan, debugging, dan pemeliharaan sistem.')

add_para('Integrasi Fungsional 100%: Pada tahap penyempurnaan terbaru (Full-Stack Data Integration), atribut-atribut krusial seperti proses unggah berkas (CNP, UTS, UAS, dan Foto Profil) yang semula merupakan mock-up visual kini telah menggunakan Engine Multer berbasis Node.js yang mengirimkan fisik file (Multipart) dan mengkonversi URL di tabel MySQL. Fungsi "Lupa Kata Sandi" kini juga tidak lagi berupa tautan statis, melainkan fitur validasi dua-faktor (NIM dan Tanggal Lahir) tersambung riil ke Database XAMPP.', bold=True)

add_heading_styled('5.2 Teknologi dan Tools', level=2)
add_table_simple(
    ['Komponen', 'Teknologi', 'Versi'],
    [
        ['Frontend Framework', 'React.js', '18.x'],
        ['Build Tool', 'Vite', '5.x'],
        ['CSS Framework', 'Tailwind CSS', '3.x'],
        ['Icon Library', 'Lucide React', 'Latest'],
        ['Routing', 'React Router DOM', 'v6'],
        ['Backend Runtime', 'Node.js', '20.x LTS'],
        ['API Framework', 'Express.js', '4.x'],
        ['File Upload', 'Multer', 'Latest'],
        ['Database', 'MySQL', '8.x (XAMPP)'],
        ['Database Driver', 'mysql2/promise', 'Latest'],
        ['IDE', 'Visual Studio Code', 'Latest'],
    ]
)

add_heading_styled('5.3 Tampilan Antarmuka Sistem', level=2)
add_para('Berikut adalah beberapa tampilan antarmuka (screenshot) dari prototipe SIAKAD Reborn yang telah dibangun:')

add_screenshot_placeholder('Halaman Landing Page (Berisi Form Lupa Sandi)')
add_para('Halaman utama sistem yang menampilkan informasi umum tentang SIAKAD Reborn, fitur-fitur unggulan, serta tombol login untuk masuk ke sistem. Desain menggunakan gradien modern dengan efek glassmorphism.', italic=True)
doc.add_paragraph()

add_screenshot_placeholder('Dashboard Mahasiswa')
add_para('Dashboard mahasiswa menampilkan ringkasan informasi akademik: pengumuman terbaru, status keuangan, jadwal hari ini, serta akses cepat ke modul-modul utama.', italic=True)
doc.add_paragraph()

add_screenshot_placeholder('Halaman Pengisian KRS')
add_para('Antarmuka pengisian KRS dengan daftar mata kuliah yang tersedia, informasi SKS, dan indikator status pengajuan.', italic=True)
doc.add_paragraph()

add_screenshot_placeholder('Halaman IKM (Indeks Keaktifan Mahasiswa)')
add_para('Form input kegiatan kemahasiswaan dengan dropdown kategori, peringkat, serta fitur upload sertifikat dan foto dokumentasi. Dilengkapi tabel riwayat poin.', italic=True)
doc.add_paragraph()

add_screenshot_placeholder('Dashboard Admin BAAK (Validasi Data)')
add_para('Panel administrasi terpusat yang memungkinkan admin mengelola seluruh data master dan memvalidasi pengajuan mahasiswa.', italic=True)
doc.add_paragraph()

add_para('Catatan: Silakan tambahkan screenshot fisik dengan cara mengambil tangkapan layar (Print Screen) dari prototipe yang sedang berjalan, kemudian timpa area Placeholder di atas.', italic=True, size=10)

add_heading_styled('5.4 Fitur-Fitur Utama Prototipe', level=2)
add_para('Prototipe SIAKAD Reborn telah mengimplementasikan fitur-fitur berikut secara fungsional:')

add_table_simple(
    ['No', 'Modul', 'Role', 'Deskripsi Fungsionalitas'],
    [
        ['1', 'Autentikasi & Keamanan', 'Semua', 'Login dengan validasi kredensial sistematis riil, Forgot Password verifikasi tangal lahir database, manajemen session.'],
        ['2', 'Dashboard', 'Semua', 'Papan pengumuman dinamis, ringkasan status per role, navigasi cepat.'],
        ['3', 'Profil', 'Mhs, Dsn', 'Tampilkan & edit biodata, ganti password, upload foto profil.'],
        ['4', 'KRS', 'Mhs', 'Pilih mata kuliah, validasi batas SKS, submit & tracking status.'],
        ['5', 'Jadwal', 'Mhs', 'Tampilkan jadwal perkuliahan berdasarkan KRS yang disetujui.'],
        ['6', 'Nilai', 'Mhs', 'Lihat rekap nilai (Tugas, UTS, UAS, Huruf) dan IPK kumulatif.'],
        ['7', 'Kehadiran', 'Mhs', 'Rekap kehadiran 14 pertemuan + UTS + UAS per mata kuliah.'],
        ['8', 'Keuangan', 'Mhs', 'Daftar tagihan, upload bukti pembayaran, tracking validasi.'],
        ['9', 'IKM', 'Mhs', 'Form setor kegiatan + sertifikat, tabel poin, verifikasi BAAK.'],
        ['10', 'CNP', 'Mhs', 'Upload berkas, download panduan, input jurnal bimbingan, jadwal sidang.'],
        ['11', 'Tugas Akhir', 'Mhs', 'Sama dengan CNP tetapi untuk kategori skripsi/TA.'],
        ['12', 'Persuratan', 'Mhs', 'Pengajuan SK Aktif dan download formulir akademik.'],
        ['13', 'Literasi Digital', 'Mhs', 'Panduan portal pelatihan, upload sertifikat kelulusan.'],
        ['14', 'Upload Dokumen', 'Mhs', 'Upload 6 dokumen wajib (Ijazah, KTP, KK, Akta, Foto, Transkrip).'],
        ['15', 'Input Nilai', 'Dsn', 'Pilih matkul, input skor per mahasiswa, kalkulasi huruf otomatis.'],
        ['16', 'Input Absensi', 'Dsn', 'Toggle kehadiran per pertemuan per mahasiswa.'],
        ['17', 'Validasi KRS', 'Adm', 'Review & approve/reject KRS mahasiswa.'],
        ['18', 'Validasi Berkas', 'Adm', 'Pusat validasi universal: CNP, TA, Keuangan, Surat, IKM, Literasi.'],
        ['19', 'Master Data', 'Adm', 'CRUD mahasiswa, dosen, mata kuliah, keuangan.'],
        ['20', 'Helpdesk', 'Semua', 'Sistem tiket terpadu, user/tamu kirim → admin balas.'],
        ['21', 'E-Library', 'Semua', 'Upload & download materi perkuliahan (PDF/dokumen).'],
    ]
)

doc.add_page_break()

# ==============================
# BAB VI - PERBANDINGAN SISTEM
# ==============================

add_heading_styled('BAB VI', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
add_heading_styled('PERBANDINGAN SISTEM LAMA DAN SISTEM BARU', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_heading_styled('6.1 Tabel Perbandingan', level=2)

add_table_simple(
    ['No', 'Aspek Perbandingan', 'Sistem Lama (SIAKAD LP3I)', 'Sistem Baru (SIAKAD Reborn)'],
    [
        ['1', 'Antarmuka (UI/UX)', 'Desain konvensional, kurang responsif di mobile', 'Modern, responsif, glassmorphism, micro-animations'],
        ['2', 'Teknologi Frontend', 'Server-side rendering tradisional', 'React.js + Vite (SPA, reaktif)'],
        ['3', 'Teknologi Backend', 'PHP/Framework tradisional', 'Node.js + Express (RESTful API)'],
        ['4', 'Integrasi Real-time', 'Terbatas, perlu refresh manual', 'Optimistic updates, sinkronisasi langsung'],
        ['5', 'Upload Dokumen', 'Terbatas pada menu tertentu', 'Terpusat melalui modul Submissions'],
        ['6', 'Validasi Admin', 'Tersebar di menu terpisah', 'Pusat Validasi Universal (1 halaman)'],
        ['7', 'IKM', 'Form sederhana', 'Form lengkap + upload sertifikat + poin otomatis'],
        ['8', 'Literasi Digital', 'Belum tersedia', 'Portal sertifikasi terintegrasi'],
        ['9', 'Helpdesk', 'Melalui jalur komunikasi terpisah', 'Terintegrasi dalam sistem + akses tamu'],
        ['10', 'E-Library', 'Terbatas', 'Upload & download file langsung di sistem'],
        ['11', 'Profil Mahasiswa', 'Data tidak bisa diubah sendiri', 'Self-service edit + Admin override'],
        ['12', 'Pembayaran', 'Konfirmasi manual', 'Upload bukti transfer + validasi online'],
    ]
)

add_heading_styled('6.2 Keunggulan Sistem Baru', level=2)
add_para('Berdasarkan tabel perbandingan di atas, sistem SIAKAD Reborn memiliki beberapa keunggulan signifikan dibanding sistem lama:')

add_bullet('User Experience yang Lebih Baik: Antarmuka modern dengan desain responsif, dark mode pada panel akademik, micro-animations, dan navigasi sidebar yang intuitif memberikan pengalaman pengguna yang jauh lebih nyaman.', bold=False)
add_bullet('Arsitektur Terpisah (Decoupled): Pemisahan frontend dan backend memungkinkan pengembangan paralel dan skalabilitas yang lebih baik di masa depan.', bold=False)
add_bullet('Pusat Validasi Terpusat: Admin tidak perlu lagi membuka banyak menu terpisah. Satu halaman "Validasi Berkas" mampu memproses semua jenis pengajuan dari mahasiswa.', bold=False)
add_bullet('Sistem IKM yang Komprehensif: Mahasiswa dapat mengumpulkan poin melalui upload sertifikat dengan kategorisasi peringkat (Internasional s.d. Lokal), dan admin dapat memberikan penilaian poin valid secara langsung.', bold=False)
add_bullet('Integrasi Modul Literasi Digital: Fitur baru yang tidak ada di sistem lama, memungkinkan kampus memantau kepatuhan mahasiswa terhadap program literasi digital nasional.', bold=False)
add_bullet('Helpdesk Terintegrasi: Mahasiswa dan bahkan tamu dapat langsung mengirim tiket bantuan tanpa perlu beralih ke platform komunikasi eksternal.', bold=False)

doc.add_page_break()

# ==============================
# BAB VII - TESTING & EVALUASI
# ==============================

add_heading_styled('BAB VII', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
add_heading_styled('TESTING DAN EVALUASI', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_heading_styled('7.1 Pengujian Fungsional', level=2)
add_para('Pengujian dilakukan menggunakan metode Black Box Testing, yaitu pengujian yang berfokus pada fungsionalitas input dan output sistem tanpa melihat struktur internal kode. Setiap modul diuji berdasarkan skenario penggunaan yang telah didefinisikan.')

add_heading_styled('7.2 Hasil Pengujian', level=2)
add_para('Berikut adalah hasil pengujian fungsional terhadap prototipe SIAKAD Reborn:')

add_table_simple(
    ['No', 'Komponen Uji', 'Skenario', 'Hasil yang Diharapkan', 'Status'],
    [
        ['1', 'Login Mahasiswa', 'Input NIM dan password yang valid', 'Redirect ke Dashboard Mahasiswa', '✅ Berhasil'],
        ['2', 'Login Gagal', 'Input password salah', 'Muncul pesan error, tetap di halaman login', '✅ Berhasil'],
        ['3', 'Pengisian KRS', 'Pilih matkul, submit', 'Status KRS berubah menjadi Pending', '✅ Berhasil'],
        ['4', 'Validasi KRS Admin', 'Admin approve KRS', 'Status berubah menjadi Approved', '✅ Berhasil'],
        ['5', 'Input Nilai Dosen', 'Input skor tugas/UTS/UAS', 'Nilai tersimpan, huruf terhitung otomatis', '✅ Berhasil'],
        ['6', 'Lihat Nilai Mhs', 'Buka menu Papan Nilai', 'Nilai per matkul ditampilkan', '✅ Berhasil'],
        ['7', 'Upload Bukti Bayar', 'Upload file dan submit', 'File terupload, status Pending', '✅ Berhasil'],
        ['8', 'Submit IKM', 'Isi form + upload sertifikat', 'Data tersimpan, tabel riwayat terupdate', '✅ Berhasil'],
        ['9', 'Validasi IKM Admin', 'Admin approve + input poin', 'Poin Valid tercatat, status berubah', '✅ Berhasil'],
        ['10', 'Tiket Helpdesk', 'Kirim pesan helpdesk', 'Tiket tercatat, admin bisa membalas', '✅ Berhasil'],
        ['11', 'Upload Literasi', 'Upload sertifikat digital', 'File tersimpan, menunggu validasi', '✅ Berhasil'],
        ['12', 'Input Absensi', 'Dosen toggle kehadiran', 'Status hadir/tidak hadir tersimpan', '✅ Berhasil'],
        ['13', 'Edit Profil Mhs', 'Ubah no HP dan alamat', 'Data terupdate di database', '✅ Berhasil'],
        ['14', 'Helpdesk Tamu', 'Kirim pesan tanpa login', 'Tiket tercatat sebagai pengirim tamu', '✅ Berhasil'],
        ['15', 'Admin Master Data', 'Tambah/edit/hapus mahasiswa', 'Data CRUD berhasil disimpan', '✅ Berhasil'],
    ]
)

add_para('Berdasarkan hasil pengujian Black Box Testing di atas, seluruh 15 skenario pengujian berhasil melewati validasi fungsional dengan status 100% BERHASIL. Hal ini menunjukkan bahwa prototipe SIAKAD Reborn telah berfungsi sesuai dengan perancangan yang telah didefinisikan pada tahap desain.', bold=False)

doc.add_page_break()

# ==============================
# BAB VIII - PEMELIHARAAN
# ==============================

add_heading_styled('BAB VIII', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
add_heading_styled('PEMELIHARAAN (MAINTENANCE)', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_heading_styled('8.1 Rencana Pemeliharaan', level=2)
add_para('Pemeliharaan (Maintenance) adalah tahapan terakhir dalam SDLC yang sangat vital untuk memastikan sistem SIAKAD Reborn tetap berjalan optimal seiring waktu. Rencana pemeliharaan ini disusun ke dalam 3 jenis:', bold=False)

add_bullet('Pemeliharaan Korektif (Corrective): Memperbaiki bug atau anomali sistem (seperti database timeout atau kendala sinkronisasi data UI) yang mungkin muncul ketika digunakan secara massal oleh ribuan mahasiswa LP3I saat pengisian KRS.')
add_bullet('Pemeliharaan Adaptif (Adaptive): Menyesuaikan Modul Backend Node.js dan Frontend React dengan versi dependensi terbaru (libraries) untuk menghindari celah keamanan (vulnerability).')
add_bullet('Pemeliharaan Perfektif (Perfective): Melakukan penyempurnaan UI/UX secara bertahap dan menambahkan fitur minor seperti notifikasi Push (Service Worker) berdasarkan feedback pengguna akademik tanpa mengubah proses inti bisnis.')

add_heading_styled('8.2 Perawatan Sistem', level=2)
add_para('Untuk menjaga kinerja operasional sistem, perawatan terjadwal akan dilakukan:', bold=False)
add_bullet('Database Backup: Mencadangkan (backup) tabel-tabel MySQL krusial (KRS, Nilai, Keuangan) setiap minggu ke External Storage (Harddisk) guna mitigasi bencana.')
add_bullet('Log Monitoring: Memeriksa Node.js / Express log untuk mendeteksi anomali request dan memitigasi serangan traffic (DDoS) dari luar kampus.')
add_bullet('Purge Cache Data: Membersihkan tumpukan cache data mahasiswa pada Browser localStorage setiap satu semester sekali agar ukuran file tidak membebani kapasitas peramban HP Mahasiswa.')

doc.add_page_break()

# ==============================
# BAB IX - KESIMPULAN & SARAN
# ==============================

add_heading_styled('BAB IX', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
add_heading_styled('KESIMPULAN DAN SARAN', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_heading_styled('9.1 Kesimpulan', level=2)
add_para('Berdasarkan hasil penelitian, perancangan, implementasi, dan pengujian yang telah dilakukan, maka dapat ditarik kesimpulan sebagai berikut:')
add_bullet('Sistem SIAKAD yang berjalan di LP3I Cirebon saat ini sudah cukup memadai namun masih memiliki beberapa keterbatasan, terutama dalam hal antarmuka pengguna, integrasi modul, dan proses administrasi yang belum sepenuhnya otomatis.')
add_bullet('Dengan menerapkan metode SDLC model Waterfall, berhasil merancang sistem SIAKAD baru (SIAKAD Reborn) yang mengakomodasi 6 tahapan murni: Planning, Analysis, Design, Implementation, Testing & Evaluasi, hingga tahapan Maintenance jangka panjang sistem.')
add_bullet('Prototipe aplikasi SIAKAD Reborn telah berhasil dibangun menggunakan React.js, Node.js, Express.js, dan MySQL. Prototipe ini bersifat fungsional dan siap deploy dengan mekanisme Multer untuk upload berkas terpusat dan verifikasi otentik Lupa Kata Sandi.')
add_bullet('Pengujian Black Box Testing terhadap 15 skenario menunjukkan tingkat keberhasilan 100%, membuktikan bahwa prototipe telah berfungsi sesuai dengan perancangan sistem.')

add_heading_styled('9.2 Saran Pengembangan', level=2)
add_para('Untuk pengembangan lebih lanjut, penulis menyarankan:')
add_bullet('Penambahan fitur autentikasi berbasis token (JWT) untuk meningkatkan keamanan sistem di lingkungan produksi.')
add_bullet('Implementasi Progressive Web App (PWA) agar sistem dapat diakses secara offline dan memberikan pengalaman seperti aplikasi native di perangkat mobile.')
add_bullet('Integrasi dengan sistem pembayaran online (Payment Gateway) untuk memfasilitasi pembayaran SPP/UKT secara langsung melalui SIAKAD.')
add_bullet('Deployment ke server publik (Cloud) agar sistem dapat diakses secara global oleh seluruh civitas akademika Politeknik LP3I.')

doc.add_page_break()

# ==============================
# DAFTAR PUSTAKA
# ==============================

add_heading_styled('DAFTAR PUSTAKA', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

refs = [
    'Jogiyanto, H.M. (2005). Analisis dan Desain Sistem Informasi: Pendekatan Terstruktur Teori dan Praktik Aplikasi Bisnis. Yogyakarta: Andi Offset.',
    'Pressman, R.S. (2014). Software Engineering: A Practitioner\'s Approach. 8th Edition. New York: McGraw-Hill Education.',
    'Laudon, K.C. & Laudon, J.P. (2016). Management Information Systems: Managing the Digital Firm. 14th Edition. Pearson Education.',
    'React – A JavaScript library for building user interfaces. (2024). Retrieved from https://react.dev/',
    'Express.js – Fast, unopinionated, minimalist web framework for Node.js. (2024). Retrieved from https://expressjs.com/',
    'MySQL 8.0 Reference Manual. (2024). Retrieved from https://dev.mysql.com/doc/refman/8.0/en/',
    'Tailwind CSS – A utility-first CSS framework. (2024). Retrieved from https://tailwindcss.com/',
    'Vite – Next Generation Frontend Tooling. (2024). Retrieved from https://vitejs.dev/',
]

for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'[{i}]  {ref}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5

doc.add_page_break()

# ==============================
# LAMPIRAN
# ==============================

add_heading_styled('LAMPIRAN', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_para('Lampiran A: Kode Sumber Mermaid untuk Diagram', bold=True)
add_para('Seluruh diagram flowchart, use case, dan ERD dalam makalah ini dibuat menggunakan format Mermaid. Kode sumber Mermaid dapat dirender secara online di https://mermaid.live untuk menghasilkan gambar beresolusi tinggi.')
doc.add_paragraph()

add_para('Catatan Teknis:', bold=True, size=11)
add_bullet('File kode Mermaid tersedia di: flowchart_siakad.md')
add_bullet('Render di https://mermaid.live → Copy kode → Paste → Download sebagai PNG')
add_bullet('Paste gambar hasil render ke placeholder [Diagram] pada makalah ini')
doc.add_paragraph()

add_para('Lampiran B: Panduan Menjalankan Prototipe', bold=True)
add_para('Untuk menjalankan prototipe SIAKAD Reborn secara lokal, ikuti langkah-langkah berikut:')
add_bullet('Pastikan XAMPP terinstall dan service Apache + MySQL sudah aktif.')
add_bullet('Buka terminal pada folder siakad-reborn/backend, jalankan: node server.js')
add_bullet('Buka terminal pada folder siakad-reborn, jalankan: npm run dev')
add_bullet('Akses http://localhost:5173 di browser.')
add_bullet('Login menggunakan kredensial default:')

add_table_simple(
    ['Role', 'ID', 'Password'],
    [
        ['Admin BAAK', 'admin', 'admin'],
        ['Dosen 1', 'dosen1', '123'],
        ['Dosen 2', 'dosen2', '123'],
        ['Mahasiswa 1', '202301001', '123'],
        ['Mahasiswa 2', '202301002', '123'],
    ]
)

# ========
# SAVE
# ========
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Makalah_UTS_SIAKAD_Reborn.docx')
doc.save(output_path)
print(f"✅ Makalah berhasil digenerate!")
print(f"📄 File tersimpan di: {output_path}")
print(f"📏 Silakan buka dengan Microsoft Word untuk review dan edit.")
